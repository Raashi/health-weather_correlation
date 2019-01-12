import os
import types
from io import BytesIO

from matplotlib.figure import Figure
from matplotlib.backends.backend_qt5 import FigureCanvasQT

from docx import Document
from docx.shared import Cm
from openpyxl import load_workbook

FACTORS = [
    "Без нагрузки",
    "С физической нагрузкой",
    "С эмоциональной нагрузкой",
    "После отдыха"
]


class XLSXParseError(ValueError):
    pass


# noinspection PyTypeChecker,PyUnresolvedReferences
class Printer:
    def __init__(self, destination, func, *args):
        self.destination = destination
        if self.destination == 'doc':
            self.doc = Document()
        elif self.destination == 'ui':
            self.doc = ''
        else:
            raise ValueError("Неизвестное назначение")
        func(*args, self)

    def print(self, destination_obj=None):
        if self.destination == 'ui':
            return self.doc
            # destination_obj.setFontPointSize(14)
            # destination_obj.insertPlainText(self.doc)
            #
            # destination_obj.installEventFilter
            # # destination_obj.resizeEvent = types.MethodType(resizeEvent, destination_obj)
            # # line_count = destination_obj.document().lineCount() + 2
            # # metrics = QFontMetrics(destination_obj.currentFont())
            # # line_spacing = metrics.lineSpacing()
            # # height = line_count * line_spacing
            # # destination_obj.setFixedHeight(height)
        else:
            save_docx(self.doc, destination_obj)
            return True

    def add_heading(self, s, size):
        if self.destination == 'doc':
            self.doc.add_heading(s, size)
        else:
            self.doc += '\n-- {} --\n'.format(s)

    def add_paragraph(self, s):
        if self.destination == 'doc':
            self.doc.add_paragraph(s)
        else:
            self.doc += s + '\n'

    def add_picture(self, pic):
        if self.destination == 'doc':
            self.doc.add_picture(pic)


def print_report(destination, func, *args):
    return Printer(destination, func, *args).print()


def is_float(s: str):
    try:
        return float(s)
    except ValueError:
        return False


def file_base_name(filename: str):
    return os.path.splitext(os.path.basename(filename))[0]


def read_sample(filename):
    """Чтение эталонов"""
    with open(filename) as file:
        data = [row.strip() for row in file]

    for idx, el in enumerate(data):
        # noinspection PyTypeChecker
        data[idx] = float(el)
    return data


def read_xlsx_sample(filename):
    wb = load_workbook(filename=filename)
    sheet = wb.get_active_sheet()

    datas = []
    for col in range(1, 4 + 1):
        data, row = [], 1
        val = sheet.cell(row, col).value
        if not val:
            datas.append(None)
            continue

        val = val.strip()
        # Пропустим первую строку, если там не число (заголовок, например)
        if is_float(val) is False:
            row += 1
        while sheet.cell(row, col).value is not None and sheet.cell(row, col).value.strip():
            try:
                data.append(float(sheet.cell(row, col).value))
            except ValueError:
                err = 'Ошибка при парсе файла {}, страницы {}, ячейки {}{}'.format(filename,
                                                                                   sheet.title,
                                                                                   'ABCD'[col],
                                                                                   row)
                print(err)
                raise XLSXParseError(err)
            row += 1
        datas.append(data)
    return datas


def patient_suffix(filename: str, suffix):
    return filename[:filename.rfind('.')] + suffix + filename[filename.rfind('.'):]


def plot_image(plot_func, *args, **kwargs):
    figure = Figure(dpi=100)
    canvas = FigureCanvasQT(figure)
    plot_func(*args, figure)

    buffer = BytesIO()
    canvas.print_figure(buffer)
    if kwargs.get('io', False):
        buffer.seek(0)
        return buffer
    else:
        return buffer.getvalue()


def create_docx():
    doc = Document()
    doc.core_properties.author = "Молчанов В.А."
    return doc


def save_docx(doc, obj):
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
    doc.save(obj)


if __name__ == '__main__':
    print('\n'.join(map(str, read_xlsx_sample("samples/1_1.xlsx"))))
